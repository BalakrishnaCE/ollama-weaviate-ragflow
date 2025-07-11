from docx import Document
from pathlib import Path
import re
from app.weaviate_client.client import get_client, create_schema
from app.ollama.client import get_embedding

DEPARTMENT_KEYWORDS = [
    "BDM", "Leasing", "Projects", "Facilities", "IT", "Sales", "Operations", "Marketing", "Finance", "Accounting"
]

def extract_department_from_path(docx_path):
    # Assumes Docs/Department/filename.docx
    parts = Path(docx_path).parts
    try:
        docs_index = parts.index("Docs")
        department = parts[docs_index + 1]
        return department
    except (ValueError, IndexError):
        return "Unknown"

def upsert_department(client, name):
    collection = client.collections.get("Department")
    # Use near_vector for vector search
    embedding = get_embedding(name)
    results = collection.query.near_vector(embedding, limit=1, return_properties=["name"])
    if not results.objects:
        return collection.data.insert({"name": name})
    return None

def upsert_sop(client, title, department, version=None, date=None):
    collection = client.collections.get("SOP")
    embedding = get_embedding(title)
    results = collection.query.near_vector(embedding, limit=1, return_properties=["title"])
    if not results.objects:
        sop_obj = {
            "title": title,
            "department": department,
            "version": version or "1.0"
        }
        if date:
            sop_obj["date"] = date
        return collection.data.insert(sop_obj)
    return None

def upsert_section(client, title, content, sop, embedding=None):
    collection = client.collections.get("Section")
    data = {
        "title": title,
        "content": content,
        "sop": sop
    }
    if embedding is not None:
        data["embedding"] = embedding
    return collection.data.insert(data)

def store_section_in_weaviate(section_data):
    client = get_client()
    # v4.x: Use collection API to insert
    # Before inserting into Weaviate, force tags to string
    section_data["tags"] = ",".join(section_data["tags"]) if isinstance(section_data["tags"], list) else str(section_data["tags"])
    try:
        client.collections.get("Section").data.insert(
            properties={
                "title": section_data["title"],
                "content": section_data["content"],
                "sop": section_data["sop"],
                "embedding": section_data["embedding"],
                "tags": section_data["tags"],
                "department": section_data["department"]
            }
        )
    except Exception as e:
        print(f"Weaviate insert error: {e}")

def extract_tags(text):
    tags = []
    for dept in DEPARTMENT_KEYWORDS:
        if dept.lower() in text.lower():
            tags.append(dept)
    return tags

def ingest_docx(docx_path):
    print(f"[DEBUG] ingest_docx called with: {docx_path}")
    doc = Document(docx_path)
    department = extract_department_from_path(docx_path)
    sop_title = Path(docx_path).stem
    sections = []
    current_section = None
    current_content = []
    section_header_pattern = re.compile(r"^([A-Z][A-Za-z0-9\s\-]+):?$")
    buffer_size = 4  # chunk size for long sections
    found_headers = False

    # Automatically extract the first non-empty paragraph as 'Overview'
    overview = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            overview = text
            break
    if overview:
        sections.append({"header": "Overview", "content": overview})

    # Main section extraction
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if section_header_pattern.match(text):
            found_headers = True
            if current_section:
                # Chunk long sections
                if len(current_content) > buffer_size:
                    for i in range(0, len(current_content), buffer_size):
                        chunk = current_content[i:i+buffer_size]
                        chunk_header = f"{current_section} (Part {i//buffer_size+1})" if len(current_content) > buffer_size else current_section
                        sections.append({"header": chunk_header, "content": "\n".join(chunk)})
                else:
                    sections.append({"header": current_section, "content": "\n".join(current_content)})
            current_section = text.rstrip(":")
            current_content = []
        else:
            current_content.append(text)
    if current_section:
        if len(current_content) > buffer_size:
            for i in range(0, len(current_content), buffer_size):
                chunk = current_content[i:i+buffer_size]
                chunk_header = f"{current_section} (Part {i//buffer_size+1})" if len(current_content) > buffer_size else current_section
                sections.append({"header": chunk_header, "content": "\n".join(chunk)})
        else:
            sections.append({"header": current_section, "content": "\n".join(current_content)})

    # If no headers found, treat every paragraph as a section
    if not found_headers:
        print("[DEBUG] No headers found, extracting every paragraph as a section.")
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append({"header": "Paragraph", "content": text})

    print(f"[DEBUG] Number of sections extracted: {len(sections)}")
    for sec in sections:
        print(f"[DEBUG] Section: {sec['header']}\n{sec['content'][:200]}\n---")
    client = get_client()
    create_schema(client)
    upsert_department(client, department)
    upsert_sop(client, sop_title, department)
    for sec in sections:
        if not sec["content"].strip():
            continue  # Skip empty sections
        print(f"[DEBUG] Storing section: {sec['header']}")
        embedding = get_embedding(sec["content"])
        tags = extract_tags(sec["content"])
        section_obj = {
            "title": sec["header"],
            "content": sec["content"],
            "sop": sop_title,
            "embedding": embedding,
            "tags": tags,
            "department": department
        }
        store_section_in_weaviate(section_obj)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 app/ingestion/docx_ingest.py <path-to-docx>")
    else:
        ingest_docx(sys.argv[1])